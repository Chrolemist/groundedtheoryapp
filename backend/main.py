from __future__ import annotations

from io import BytesIO
import os
from pathlib import Path
import json
import logging
import hashlib
import random
import time
import secrets
from typing import Dict, List, Optional
from uuid import uuid4

import pandas as pd
import firebase_admin
from firebase_admin import credentials, firestore
from docx import Document
from dotenv import load_dotenv
from fastapi import Body, FastAPI, File, HTTPException, Request, UploadFile, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field


def monotonic_updated_at_ms(previous: int) -> int:
    """Generate a monotonic ms timestamp based on server time.

    We use this as the canonical updated_at for project_raw to avoid client clock skew
    breaking live sync ("stale" rejects).
    """
    now_ms = int(time.time() * 1000)
    try:
        prev_ms = int(previous) if previous else 0
    except Exception:
        prev_ms = 0
    # Ensure strictly increasing when multiple updates happen quickly.
    return max(now_ms, prev_ms + 1)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
env_path = Path(__file__).resolve().parent / ".env"
load_dotenv(dotenv_path=env_path, override=False)
BUILD_TAG = os.getenv("BUILD_TAG", "local-dev")
logger.info("Build tag: %s", BUILD_TAG)
MAX_PROJECT_BYTES = int(os.getenv("MAX_PROJECT_BYTES", "900000"))
MAX_PROJECT_TOTAL_BYTES = int(os.getenv("MAX_PROJECT_TOTAL_BYTES", "900000000"))
EMPTY_OVERWRITE_MAX_BYTES = int(os.getenv("EMPTY_OVERWRITE_MAX_BYTES", "800"))
NONEMPTY_PROJECT_MIN_BYTES = int(os.getenv("NONEMPTY_PROJECT_MIN_BYTES", "1600"))
logger.info(
    "Firestore config: project_id=%s collection=%s",
    os.getenv("FIREBASE_PROJECT_ID"),
    os.getenv("FIRESTORE_COLLECTION", "projects"),
)

DEFAULT_PROJECT_ID = "default"


class DocumentItem(BaseModel):
    id: str
    title: str
    content: str
    html: str = ''
    text: str = ''


class CodeItem(BaseModel):
    id: str
    name: str
    color: str


class HighlightItem(BaseModel):
    id: str
    document_id: str
    start_index: int
    end_index: int
    code_id: str


class CategoryItem(BaseModel):
    id: str
    name: str
    contained_code_ids: List[str] = Field(default_factory=list)
    precondition: str = ''
    action: str = ''
    consequence: str = ''


class ProjectState(BaseModel):
    documents: List[DocumentItem] = Field(default_factory=list)
    codes: List[CodeItem] = Field(default_factory=list)
    highlights: List[HighlightItem] = Field(default_factory=list)
    categories: List[CategoryItem] = Field(default_factory=list)
    core_category_id: Optional[str] = None
    theory_description: str = ''


class AdminLoginPayload(BaseModel):
    password: str


ADMIN_LOCKOUT_SCHEDULE_SECONDS = [
    60,
    300,
    3600,
    86400,
    604800,
    2592000,
    31536000,
]
ADMIN_LOCKOUT_LABELS = {
    60: "1 minute",
    300: "5 minutes",
    3600: "1 hour",
    86400: "24 hours",
    604800: "1 week",
    2592000: "1 month",
    31536000: "1 year",
}
ADMIN_ATTEMPT_LIMIT = 3
admin_lockouts: Dict[str, Dict[str, float | int]] = {}
ADMIN_TOKEN_TTL_SECONDS = int(os.getenv("ADMIN_TOKEN_TTL_SECONDS", "28800"))
admin_tokens: Dict[str, float] = {}


def get_client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    if request.client:
        return request.client.host
    return "unknown"


def format_lockout_wait(seconds: int, fallback_seconds: int) -> str:
    label = ADMIN_LOCKOUT_LABELS.get(fallback_seconds)
    if label:
        return label
    if seconds < 60:
        return f"{max(1, seconds)} seconds"
    if seconds < 3600:
        minutes = max(1, int((seconds + 59) / 60))
        return f"{minutes} minutes"
    if seconds < 86400:
        hours = max(1, int((seconds + 3599) / 3600))
        return f"{hours} hours"
    days = max(1, int((seconds + 86399) / 86400))
    return f"{days} days"


def prune_admin_tokens(now: float) -> None:
    expired = [token for token, expires_at in admin_tokens.items() if expires_at <= now]
    for token in expired:
        admin_tokens.pop(token, None)


def issue_admin_token() -> str:
    token = secrets.token_urlsafe(32)
    admin_tokens[token] = time.time() + ADMIN_TOKEN_TTL_SECONDS
    return token


def validate_admin_token(token: str) -> bool:
    if not token:
        return False
    now = time.time()
    prune_admin_tokens(now)
    expires_at = admin_tokens.get(token)
    if not expires_at:
        return False
    if expires_at <= now:
        admin_tokens.pop(token, None)
        return False
    return True


def require_admin(request: Request) -> None:
    auth_header = request.headers.get("authorization") or ""
    parts = auth_header.split(" ", 1)
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(status_code=401, detail="Admin authorization required")
    token = parts[1].strip()
    if not validate_admin_token(token):
        raise HTTPException(status_code=401, detail="Admin authorization required")


class ConnectionManager:
    def __init__(self) -> None:
        self.active_connections: Dict[str, List[WebSocket]] = {}
        self.connection_users: Dict[WebSocket, str] = {}
        self.users: Dict[str, Dict[str, Dict[str, str]]] = {}
        self.client_users: Dict[str, Dict[str, str]] = {}
        self.user_connections: Dict[str, Dict[str, WebSocket]] = {}
        self.connection_projects: Dict[WebSocket, str] = {}

    def _generate_user_name(self) -> str:
        first_parts = [
            "Busiga",
            "Snälla",
            "Tokiga",
            "Glada",
            "Luriga",
            "Pigga",
            "Coola",
            "Snabba",
            "Mjuka",
            "Kloka",
            "Sköna",
            "Modiga",
        ]
        second_parts = [
            "Räven",
            "Ugglan",
            "Igelkotten",
            "Bävern",
            "Lamasen",
            "Kängurun",
            "Pandan",
            "Humlan",
            "Sälen",
            "Älgen",
            "Vesslan",
            "Grodan",
        ]
        existing: set[str] = set()
        for project_users in self.users.values():
            if not isinstance(project_users, dict):
                continue
            for user in project_users.values():
                if isinstance(user, dict):
                    name = user.get("name")
                    if isinstance(name, str) and name:
                        existing.add(name)
        max_attempts = max(20, len(first_parts) * len(second_parts))
        for _ in range(max_attempts):
            base = f"{random.choice(first_parts)} {random.choice(second_parts)}"
            if base not in existing:
                return base
        suffix = 2
        base = f"{first_parts[0]} {second_parts[0]}"
        while f"{base} {suffix}" in existing:
            suffix += 1
        return f"{base} {suffix}"

    def _generate_user_color(self, project_id: str) -> str:
        palette = [
            "#2563EB",
            "#7C3AED",
            "#DC2626",
            "#16A34A",
            "#D97706",
            "#0EA5E9",
            "#DB2777",
            "#0F766E",
        ]
        return palette[len(self.users.get(project_id, {})) % len(palette)]

    async def connect(
        self,
        websocket: WebSocket,
        client_id: Optional[str] = None,
        project_id: str = DEFAULT_PROJECT_ID,
    ) -> Dict[str, str]:
        await websocket.accept()
        user_id: Optional[str] = None
        user: Optional[Dict[str, str]] = None

        project_users = self.users.setdefault(project_id, {})
        project_client_users = self.client_users.setdefault(project_id, {})
        project_user_connections = self.user_connections.setdefault(project_id, {})
        project_connections = self.active_connections.setdefault(project_id, [])

        if client_id:
            existing_user_id = project_client_users.get(client_id)
            if existing_user_id:
                existing_user = project_users.get(existing_user_id)
                old_socket = project_user_connections.get(existing_user_id)
                if old_socket:
                    if old_socket in project_connections:
                        project_connections.remove(old_socket)
                    self.connection_users.pop(old_socket, None)
                    self.connection_projects.pop(old_socket, None)
                    try:
                        logger.info(
                            "Closing previous socket for client_id=%s user_id=%s",
                            client_id,
                            existing_user_id,
                        )
                        await old_socket.close(code=1000, reason="replaced")
                    except Exception:
                        pass
                if existing_user:
                    user_id = existing_user_id
                    user = existing_user

        if not user_id or not user:
            user_id = str(uuid4())
            user = {
                "id": user_id,
                "name": self._generate_user_name(),
                "color": self._generate_user_color(project_id),
            }
        else:
            if "name" not in user or not user.get("name"):
                user["name"] = self._generate_user_name()
            if "color" not in user or not user.get("color"):
                user["color"] = self._generate_user_color(project_id)

        project_connections.append(websocket)
        self.connection_users[websocket] = user_id
        project_users[user_id] = user
        project_user_connections[user_id] = websocket
        self.connection_projects[websocket] = project_id
        if client_id:
            project_client_users[client_id] = user_id
            project_users[user_id]["client_id"] = client_id
        return user

    def disconnect(self, websocket: WebSocket) -> None:
        project_id = self.connection_projects.pop(websocket, DEFAULT_PROJECT_ID)
        project_connections = self.active_connections.get(project_id, [])
        if websocket in project_connections:
            project_connections.remove(websocket)
        user_id = self.connection_users.pop(websocket, None)
        project_users = self.users.get(project_id, {})
        project_client_users = self.client_users.get(project_id, {})
        project_user_connections = self.user_connections.get(project_id, {})
        if user_id and user_id in project_users:
            client_id = project_users[user_id].get("client_id")
            project_users.pop(user_id, None)
            project_user_connections.pop(user_id, None)
            if client_id:
                existing = project_client_users.get(client_id)
                if existing == user_id:
                    project_client_users.pop(client_id, None)

    def get_users(self, project_id: str) -> List[Dict[str, str]]:
        return list(self.users.get(project_id, {}).values())

    async def broadcast(self, project_id: str, message: Dict) -> None:
        # Include project_id so clients can ignore stale cross-project messages.
        if isinstance(message, dict) and "project_id" not in message:
            message = {**message, "project_id": project_id}
        for connection in list(self.active_connections.get(project_id, [])):
            try:
                await connection.send_json(message)
            except Exception:
                self.disconnect(connection)

    async def broadcast_except(self, project_id: str, message: Dict, skip: WebSocket) -> None:
        # Include project_id so clients can ignore stale cross-project messages.
        if isinstance(message, dict) and "project_id" not in message:
            message = {**message, "project_id": project_id}
        for connection in list(self.active_connections.get(project_id, [])):
            if connection == skip:
                continue
            try:
                await connection.send_json(message)
            except Exception:
                self.disconnect(connection)


app = FastAPI()
manager = ConnectionManager()

DEFAULT_PROJECT_ID = "default"
in_memory_projects: Dict[str, ProjectState] = {}
in_memory_project_raws: Dict[str, Dict] = {}
yjs_update_logs: Dict[str, List[str]] = {}
yjs_update_limit = 200

firestore_client: Optional[firestore.Client] = None
last_saved_project_hash: Dict[str, str] = {}


def resolve_project_id(project_id: Optional[str]) -> str:
    return project_id or DEFAULT_PROJECT_ID


def get_in_memory_project(project_id: str) -> ProjectState:
    return in_memory_projects.setdefault(project_id, ProjectState())


def get_in_memory_project_raw(project_id: str) -> Dict:
    return in_memory_project_raws.get(project_id, {})


def set_in_memory_project(project_id: str, project_raw: Dict) -> None:
    in_memory_project_raws[project_id] = project_raw
    in_memory_projects[project_id] = normalize_project_state(project_raw)


def ensure_project_loaded(project_id: str) -> None:
    project_raw = load_project_from_firestore(None if project_id == DEFAULT_PROJECT_ID else project_id)
    if project_raw:
        set_in_memory_project(project_id, project_raw)
        project_hash = compute_project_hash(project_raw)
        if project_hash:
            last_saved_project_hash[project_id] = project_hash
        return

    if project_id not in in_memory_projects:
        in_memory_projects[project_id] = ProjectState()
    if project_id not in in_memory_project_raws and in_memory_projects[project_id].documents:
        in_memory_project_raws[project_id] = in_memory_projects[project_id].model_dump()
        project_hash = compute_project_hash(in_memory_project_raws[project_id])
        if project_hash:
            last_saved_project_hash[project_id] = project_hash


def get_firestore_client() -> Optional[firestore.Client]:
    global firestore_client
    if firestore_client is not None:
        return firestore_client

    project_id = (os.getenv("FIREBASE_PROJECT_ID") or "").strip() or None
    credentials_json = os.getenv("FIREBASE_CREDENTIALS_JSON")
    credentials_path = os.getenv("FIREBASE_CREDENTIALS_PATH")

    try:
        if not firebase_admin._apps:
            if credentials_json:
                creds_dict = json.loads(credentials_json)
                resolved_project_id = project_id or creds_dict.get("project_id")
                options = {"projectId": resolved_project_id} if resolved_project_id else {}
                firebase_admin.initialize_app(credentials.Certificate(creds_dict), options)
            elif credentials_path:
                options = {"projectId": project_id} if project_id else {}
                firebase_admin.initialize_app(credentials.Certificate(credentials_path), options)
            else:
                # Use Application Default Credentials (Cloud Run service account)
                options = {"projectId": project_id} if project_id else {}
                firebase_admin.initialize_app(credentials.ApplicationDefault(), options)

        firestore_client = firestore.client()
        logger.info(
            "Firestore client initialized. app_project_id=%s client_project=%s",
            project_id,
            getattr(firestore_client, "project", None),
        )
        return firestore_client
    except Exception as exc:
        logger.exception("Firestore not configured or failed to initialize")
        return None


def get_firestore_doc_ref():
    client = get_firestore_client()
    if not client:
        return None
    collection = os.getenv("FIRESTORE_COLLECTION", "projects")
    doc_id = os.getenv("FIRESTORE_DOC", "default")
    return client.collection(collection).document(doc_id)


def get_project_doc_ref(project_id: str):
    client = get_firestore_client()
    if not client:
        return None
    collection = os.getenv("FIRESTORE_COLLECTION", "projects")
    return client.collection(collection).document(project_id)


def compute_project_hash(project_raw: Dict) -> Optional[str]:
    try:
        if isinstance(project_raw, dict) and "updated_at" in project_raw:
            sanitized = dict(project_raw)
            sanitized.pop("updated_at", None)
        else:
            sanitized = project_raw
        payload = json.dumps(sanitized, sort_keys=True, separators=(",", ":"))
    except Exception:
        return None
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def estimate_project_bytes(project_raw: Dict) -> Optional[int]:
    try:
        payload = json.dumps(project_raw, sort_keys=True, separators=(",", ":"))
    except Exception:
        return None
    return len(payload.encode("utf-8"))


def validate_project_limits(project_raw: Dict) -> Optional[Dict[str, str]]:
    size_bytes = estimate_project_bytes(project_raw)
    if size_bytes is None:
        return {
            "reason": "project_invalid",
            "message": "Project payload could not be serialized.",
        }
    if MAX_PROJECT_BYTES > 0 and size_bytes > MAX_PROJECT_BYTES:
        return {
            "reason": "project_too_large",
            "message": (
                f"Project size {size_bytes} bytes exceeds limit {MAX_PROJECT_BYTES} bytes."
            ),
        }
    if MAX_PROJECT_TOTAL_BYTES > 0 and size_bytes > MAX_PROJECT_TOTAL_BYTES:
        return {
            "reason": "project_total_limit_reached",
            "message": (
                f"Project total size {size_bytes} bytes exceeds limit {MAX_PROJECT_TOTAL_BYTES} bytes."
            ),
        }
    return None


def is_project_effectively_empty(project_raw: Dict) -> bool:
    try:
        if not project_raw:
            return True
        documents = project_raw.get("documents")
        codes = project_raw.get("codes")
        categories = project_raw.get("categories")
        memos = project_raw.get("memos")
        core_category_id = project_raw.get("coreCategoryId") or project_raw.get("core_category_id")
        theory_html = project_raw.get("theoryHtml") or project_raw.get("theory_description")
        updated_at = project_raw.get("updated_at")

        docs_count = len(documents) if isinstance(documents, list) else 0
        codes_count = len(codes) if isinstance(codes, list) else 0
        categories_count = len(categories) if isinstance(categories, list) else 0
        memos_count = len(memos) if isinstance(memos, list) else 0

        # Treat payload as empty when it only contains metadata/updated_at and no real content.
        return (
            docs_count == 0
            and codes_count == 0
            and categories_count == 0
            and memos_count == 0
            and (not str(core_category_id or "").strip())
            and (not str(theory_html or "").strip())
            and updated_at is not None
        )
    except Exception:
        return False


def estimate_document_content_chars(project_raw: Dict) -> int:
    """Heuristic: count non-whitespace characters across document content fields."""
    try:
        documents = project_raw.get("documents")
        if not isinstance(documents, list):
            return 0
        total = 0
        for doc in documents:
            if not isinstance(doc, dict):
                continue
            # Prefer explicit text/content; HTML can be present but empty like '<p></p>'.
            text = doc.get("text")
            if not isinstance(text, str) or not text.strip():
                text = doc.get("content")
            if isinstance(text, str) and text.strip():
                total += len("".join(text.split()))
                continue
            html = doc.get("html")
            if isinstance(html, str) and html.strip():
                # Very rough HTML-to-text stripping; still better than treating any tag as content.
                stripped = "".join(html.replace("&nbsp;", " ").split())
                # Remove common empty-paragraph patterns.
                if stripped in ("<p></p>", "<p><br/></p>", "<p><br></p>"):
                    continue
                # If it contains any letter/digit, count it.
                if any(ch.isalnum() for ch in stripped):
                    total += sum(1 for ch in stripped if not ch.isspace())
        return total
    except Exception:
        return 0


def save_project_to_firestore(project_raw: Dict, project_id: Optional[str] = None) -> bool:
    global last_saved_project_hash
    doc_ref = get_project_doc_ref(project_id) if project_id else get_firestore_doc_ref()
    if not doc_ref:
        logger.error(
            "Firestore save skipped (no doc ref). project_id=%s",
            project_id or "default",
        )
        return False
    current_hash = compute_project_hash(project_raw)
    hash_key = project_id or "default"
    if current_hash and current_hash == last_saved_project_hash.get(hash_key):
        return True
    try:
        doc_ref.set(
            {
                "project": project_raw,
                "updated_at": firestore.SERVER_TIMESTAMP,
            },
            merge=True,
        )
        if current_hash:
            last_saved_project_hash[hash_key] = current_hash
        try:
            logger.info(
                "Saved project to Firestore. project_id=%s size_bytes=%s",
                project_id or "default",
                estimate_project_bytes(project_raw),
            )
        except Exception:
            pass
        return True
    except Exception as exc:
        logger.exception(
            "Failed to save project to Firestore. project_id=%s",
            project_id or "default",
        )
        return False


def load_project_from_firestore(project_id: Optional[str] = None) -> Optional[Dict]:
    doc_ref = get_project_doc_ref(project_id) if project_id else get_firestore_doc_ref()
    if not doc_ref:
        return None
    try:
        snapshot = doc_ref.get()
        if not snapshot.exists:
            return None
        data = snapshot.to_dict() or {}
        return data.get("project")
    except Exception as exc:
        logger.exception(
            "Failed to load project from Firestore. project_id=%s",
            project_id or "default",
        )
        return None


def serialize_project_summary(snapshot) -> Dict[str, Optional[str]]:
    data = snapshot.to_dict() or {}
    name = data.get("name") or "Untitled project"
    updated_at = data.get("updated_at")
    created_at = data.get("created_at")
    def to_iso(value):
        if hasattr(value, "isoformat"):
            return value.isoformat()
        return None
    return {
        "id": snapshot.id,
        "name": name,
        "updated_at": to_iso(updated_at),
        "created_at": to_iso(created_at),
    }

allowed_origins = [origin.strip() for origin in os.getenv("ALLOWED_ORIGINS", "http://localhost:5173").split(",") if origin.strip()]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def load_persisted_project() -> None:
    ensure_project_loaded(DEFAULT_PROJECT_ID)
    if get_in_memory_project_raw(DEFAULT_PROJECT_ID):
        logger.info("Loaded project from Firestore")

FRONTEND_DIST = Path(__file__).resolve().parent.parent / "frontend" / "dist"

# Mount assets explicitly to avoid intercepting root requests (and WebSockets)
if (FRONTEND_DIST / "assets").exists():
    app.mount("/assets", StaticFiles(directory=FRONTEND_DIST / "assets"), name="assets")


def get_code_by_id(code_id: str) -> Optional[CodeItem]:
    project = get_in_memory_project(DEFAULT_PROJECT_ID)
    return next((code for code in project.codes if code.id == code_id), None)


def get_category_by_id(category_id: str) -> Optional[CategoryItem]:
    project = get_in_memory_project(DEFAULT_PROJECT_ID)
    return next((category for category in project.categories if category.id == category_id), None)


def get_document_by_id(document_id: str) -> Optional[DocumentItem]:
    project = get_in_memory_project(DEFAULT_PROJECT_ID)
    return next((document for document in project.documents if document.id == document_id), None)


def normalize_project_state(payload: Dict) -> ProjectState:
    documents = []
    for doc in payload.get("documents", []):
        content = doc.get("content") or doc.get("text") or ""
        documents.append(
            DocumentItem(
                id=doc.get("id", ""),
                title=doc.get("title", ""),
                content=content,
                html=doc.get("html", ""),
                text=doc.get("text", ""),
            )
        )

    categories = []
    for category in payload.get("categories", []):
        categories.append(
            CategoryItem(
                id=category.get("id", ""),
                name=category.get("name", ""),
                contained_code_ids=category.get("contained_code_ids", [])
                or category.get("codeIds", []),
                precondition=category.get("precondition", ""),
                action=category.get("action", ""),
                consequence=category.get("consequence", ""),
            )
        )

    return ProjectState(
        documents=documents,
        codes=[
            CodeItem(
                id=code.get("id", ""),
                name=code.get("name") or code.get("label") or "",
                color=code.get("color") or code.get("colorHex") or "#E2E8F0",
            )
            for code in payload.get("codes", [])
        ],
        highlights=[
            HighlightItem(
                id=highlight.get("id", ""),
                document_id=highlight.get("document_id", ""),
                start_index=highlight.get("start_index", 0),
                end_index=highlight.get("end_index", 0),
                code_id=highlight.get("code_id", ""),
            )
            for highlight in payload.get("highlights", [])
        ],
        categories=categories,
        core_category_id=payload.get("core_category_id")
        or payload.get("coreCategoryId"),
        theory_description=payload.get("theory_description", "")
        or payload.get("theoryHtml", ""),
    )


@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket) -> None:
    logger.info("New WebSocket connection request")
    try:
        client_id = websocket.query_params.get("client_id")
        project_id = resolve_project_id(websocket.query_params.get("project_id"))
        user = await manager.connect(websocket, client_id, project_id)
        user_agent = websocket.headers.get("user-agent")
        origin = websocket.headers.get("origin")
        user_id = user.get("id", "unknown") if isinstance(user, dict) else "unknown"
        user_name = user.get("name", "unknown") if isinstance(user, dict) else "unknown"
        if isinstance(user, dict):
            if not user.get("name"):
                user["name"] = user_name
            if not user.get("color"):
                user["color"] = "#7C3AED"
        logger.info(
            "WebSocket accepted. User: %s (%s) client_id=%s remote=%s ua=%s origin=%s",
            user_id,
            user_name,
            client_id,
            websocket.client,
            user_agent,
            origin,
        )
    except Exception as e:
        logger.exception("WebSocket connection failed")
        try:
            await websocket.close(code=1011)
        except Exception:
            pass
        return

    ensure_project_loaded(project_id)
    project_state = get_in_memory_project(project_id)
    project_raw = get_in_memory_project_raw(project_id)

    # Ensure in-memory project_raw carries a canonical monotonic updated_at.
    if isinstance(project_raw, dict):
        prev_updated_at = project_raw.get("updated_at")
        prev_ms = prev_updated_at if isinstance(prev_updated_at, int) else 0
        next_ms = max(int(time.time() * 1000), prev_ms)
        if next_ms and prev_ms != next_ms:
            project_raw["updated_at"] = next_ms

    await websocket.send_json(
        {
            "type": "hello",
            "project_id": project_id,
            "user": user,
            "users": manager.get_users(project_id),
            "project": project_state.model_dump(),
            "project_raw": project_raw or None,
        }
    )
    logger.info("WS hello sent: project_id=%s user_id=%s users=%s", project_id, user.get("id"), len(manager.get_users(project_id)))
    project_yjs_log = yjs_update_logs.get(project_id, [])
    # Always send a sync message (even if empty) so clients can complete their
    # startup handshake and seed initial content when needed.
    await websocket.send_json({"type": "yjs:sync", "project_id": project_id, "updates": project_yjs_log})
    await manager.broadcast(
        project_id,
        {"type": "presence:update", "users": manager.get_users(project_id)},
    )
    try:
        while True:
            raw = await websocket.receive_text()
            # logger.debug(f"Received message: {raw[:100]}...")  # Optional: Log incoming messages
            try:
                data = json.loads(raw)
            except Exception:
                continue

            message_type = data.get("type")
            user_id = manager.connection_users.get(websocket)
            if message_type:
                logger.info("WS message: project_id=%s user_id=%s type=%s", project_id, user_id, message_type)
            if message_type == "ping":
                await websocket.send_json({"type": "pong", "ts": data.get("ts")})
                continue
            if message_type == "presence:rename" and user_id:
                next_name = str(data.get("name", "")).strip()
                if next_name:
                    project_users = manager.users.get(project_id, {})
                    if user_id in project_users:
                        project_users[user_id]["name"] = next_name[:40]
                    await manager.broadcast(
                        project_id,
                        {"type": "presence:update", "users": manager.get_users(project_id)},
                    )
            elif message_type == "cursor:update" and user_id:
                payload = {
                    "type": "cursor:update",
                    "userId": user_id,
                    "cursor": data.get("cursor"),
                }
                await manager.broadcast_except(project_id, payload, websocket)
            elif message_type == "cursor:clear" and user_id:
                await manager.broadcast_except(
                    project_id,
                    {"type": "cursor:clear", "userId": user_id},
                    websocket,
                )
            elif message_type == "selection:update" and user_id:
                payload = {
                    "type": "selection:update",
                    "userId": user_id,
                    "selection": data.get("selection"),
                }
                await manager.broadcast_except(project_id, payload, websocket)
            elif message_type == "selection:clear" and user_id:
                await manager.broadcast_except(
                    project_id,
                    {"type": "selection:clear", "userId": user_id},
                    websocket,
                )
            elif message_type == "yjs:update" and user_id:
                update = data.get("update")
                if update:
                    project_yjs_log = yjs_update_logs.setdefault(project_id, [])
                    project_yjs_log.append(update)
                    if len(project_yjs_log) > yjs_update_limit:
                        project_yjs_log[:] = project_yjs_log[-yjs_update_limit:]
                    await manager.broadcast_except(
                        project_id,
                        {
                            "type": "yjs:update",
                            "update": update,
                        },
                        websocket,
                    )
            elif message_type == "project:update" and user_id:
                project_payload = data.get("project", {})
                project_raw = data.get("project_raw") or project_payload
                if isinstance(project_raw, dict):
                    limit_error = validate_project_limits(project_raw)
                    if limit_error:
                        await websocket.send_json(
                            {
                                "type": "project:save:error",
                                **limit_error,
                            }
                        )
                        continue

                    # Safety net: avoid wiping a real project due to a client race/bug.
                    # The REST save endpoint has similar guardrails; apply them here too
                    # because `project:update` can persist to Firestore.
                    storage_id = None if project_id == DEFAULT_PROJECT_ID else project_id
                    existing = load_project_from_firestore(storage_id)
                    existing_doc_chars = (
                        estimate_document_content_chars(existing)
                        if isinstance(existing, dict)
                        else 0
                    )
                    incoming_doc_chars = estimate_document_content_chars(project_raw)
                    looks_like_content_wipe = (
                        isinstance(existing, dict)
                        and existing_doc_chars > 0
                        and incoming_doc_chars == 0
                    )
                    if looks_like_content_wipe:
                        logger.warning(
                            "Refusing content wipe via WS. project_id=%s incoming_doc_chars=%s existing_doc_chars=%s",
                            project_id,
                            incoming_doc_chars,
                            existing_doc_chars,
                        )
                        await websocket.send_json(
                            {
                                "type": "project:save:error",
                                "reason": "content_wipe",
                                "message": "Refusing to overwrite a project with a payload that contains no document content.",
                            }
                        )
                        continue
                # Canonicalize updated_at using server time to avoid client clock skew.
                current_raw = get_in_memory_project_raw(project_id)
                current_updated_at = (
                    current_raw.get("updated_at")
                    if isinstance(current_raw, dict)
                    else 0
                )
                previous_ms = current_updated_at if isinstance(current_updated_at, int) else 0
                project_raw["updated_at"] = monotonic_updated_at_ms(previous_ms)
                set_in_memory_project(project_id, project_raw)
                saved_ok = save_project_to_firestore(project_raw, project_id=storage_id)
                if not saved_ok:
                    await websocket.send_json(
                        {
                            "type": "project:save:error",
                            "reason": "firestore_error",
                            "message": "Failed to save project to Firestore.",
                        }
                    )
                await manager.broadcast(
                    project_id,
                    {
                        "type": "project:update",
                        "project": get_in_memory_project(project_id).model_dump(),
                        "project_raw": get_in_memory_project_raw(project_id),
                        "sender_id": user_id,
                    },
                )
    except WebSocketDisconnect as exc:
        user_agent = websocket.headers.get("user-agent")
        origin = websocket.headers.get("origin")
        logger.info(
            "WebSocket disconnected: %s (code=%s) client_id=%s remote=%s ua=%s origin=%s",
            manager.connection_users.get(websocket),
            getattr(exc, "code", None),
            client_id,
            websocket.client,
            user_agent,
            origin,
        )
        user_id = manager.connection_users.get(websocket)
        manager.disconnect(websocket)
        await manager.broadcast(
            project_id,
            {"type": "presence:update", "users": manager.get_users(project_id)},
        )
        if user_id:
            await manager.broadcast(project_id, {"type": "cursor:clear", "userId": user_id})
            await manager.broadcast(
                project_id,
                {"type": "selection:clear", "userId": user_id},
            )


@app.get("/project/state")
async def get_project_state() -> Dict[str, Optional[Dict]]:
    project_id = DEFAULT_PROJECT_ID
    ensure_project_loaded(project_id)
    project_raw = get_in_memory_project_raw(project_id)
    project_state = get_in_memory_project(project_id)
    return {
        "project_raw": project_raw or None,
        "project": project_state.model_dump(),
    }


@app.get("/projects")
async def list_projects() -> Dict[str, List[Dict[str, Optional[str]]]]:
    client = get_firestore_client()
    if not client:
        return {"projects": []}
    collection = os.getenv("FIRESTORE_COLLECTION", "projects")
    try:
        query = client.collection(collection).order_by(
            "updated_at", direction=firestore.Query.DESCENDING
        )
        snapshots = query.stream()
        return {"projects": [serialize_project_summary(doc) for doc in snapshots]}
    except Exception as exc:
        logger.warning(f"Failed to list projects: {exc}")
        return {"projects": []}


@app.post("/admin/login")
async def admin_login(payload: AdminLoginPayload, request: Request) -> Dict[str, object]:
    expected = (os.getenv("ADMIN_PASSWORD") or "").strip()
    if not expected:
        return {"ok": False, "message": "Admin password not configured"}
    client_key = get_client_ip(request)
    state = admin_lockouts.setdefault(
        client_key,
        {"failures": 0, "lockout_until": 0, "step": -1},
    )
    now = time.time()
    lockout_until = float(state.get("lockout_until", 0))
    if lockout_until > now:
        remaining = int(lockout_until - now)
        step_index = int(state.get("step", -1))
        step_index = max(0, min(step_index, len(ADMIN_LOCKOUT_SCHEDULE_SECONDS) - 1))
        step_duration = ADMIN_LOCKOUT_SCHEDULE_SECONDS[step_index]
        return {
            "ok": False,
            "message": (
                "Too many attempts. Try again in "
                f"{format_lockout_wait(remaining, step_duration)}."
            ),
            "retry_after_seconds": remaining,
        }
    if payload.password == expected:
        admin_lockouts.pop(client_key, None)
        token = issue_admin_token()
        return {
            "ok": True,
            "token": token,
            "expires_in_seconds": ADMIN_TOKEN_TTL_SECONDS,
        }
    state["failures"] = int(state.get("failures", 0)) + 1
    if state["failures"] % ADMIN_ATTEMPT_LIMIT == 0:
        next_step = min(
            int(state.get("step", -1)) + 1,
            len(ADMIN_LOCKOUT_SCHEDULE_SECONDS) - 1,
        )
        duration = ADMIN_LOCKOUT_SCHEDULE_SECONDS[next_step]
        state["step"] = next_step
        state["lockout_until"] = now + duration
        return {
            "ok": False,
            "message": f"Too many attempts. Try again in {format_lockout_wait(duration, duration)}.",
            "retry_after_seconds": int(duration),
        }
    return {"ok": False, "message": "Invalid password"}


@app.get("/projects/storage")
async def get_projects_storage() -> Dict[str, int]:
    client = get_firestore_client()
    if not client:
        return {
            "total_bytes": 0,
            "total_limit_bytes": MAX_PROJECT_TOTAL_BYTES,
            "project_limit_bytes": MAX_PROJECT_BYTES,
        }
    collection = os.getenv("FIRESTORE_COLLECTION", "projects")
    total_bytes = 0
    try:
        snapshots = client.collection(collection).stream()
        for snapshot in snapshots:
            data = snapshot.to_dict() or {}
            project_raw = data.get("project")
            if isinstance(project_raw, dict):
                size_bytes = estimate_project_bytes(project_raw)
                if isinstance(size_bytes, int):
                    total_bytes += size_bytes
    except Exception as exc:
        logger.warning(f"Failed to compute project storage: {exc}")
    return {
        "total_bytes": total_bytes,
        "total_limit_bytes": MAX_PROJECT_TOTAL_BYTES,
        "project_limit_bytes": MAX_PROJECT_BYTES,
    }


@app.post("/projects/purge")
async def purge_projects(request: Request) -> Dict[str, int | str]:
    require_admin(request)
    global last_saved_project_hash
    client = get_firestore_client()
    if not client:
        return {"status": "error", "message": "Firestore not available"}
    collection = os.getenv("FIRESTORE_COLLECTION", "projects")
    deleted = 0
    batch = client.batch()
    batch_count = 0
    try:
        snapshots = client.collection(collection).stream()
        for snapshot in snapshots:
            batch.delete(snapshot.reference)
            deleted += 1
            batch_count += 1
            if batch_count >= 450:
                batch.commit()
                batch = client.batch()
                batch_count = 0
        if batch_count:
            batch.commit()
        last_saved_project_hash.clear()
        return {"status": "ok", "deleted": deleted}
    except Exception as exc:
        logger.warning(f"Failed to purge projects: {exc}")
        return {"status": "error", "message": "Failed to purge projects"}


@app.post("/projects")
async def create_project(payload: Dict = Body(...)) -> Dict[str, object]:
    client = get_firestore_client()
    if not client:
        return {"status": "error", "message": "Firestore not available"}
    name = str(payload.get("name") or "New project").strip() or "New project"
    project_raw = payload.get("project_raw")
    if not isinstance(project_raw, dict):
        project_raw = {
            "documents": [],
            "codes": [],
            "categories": [],
            "memos": [],
            "coreCategoryId": "",
            "theoryHtml": "",
        }
    limit_error = validate_project_limits(project_raw)
    if limit_error:
        return {
            "status": "error",
            **limit_error,
        }
    project_id = str(uuid4())
    doc_ref = get_project_doc_ref(project_id)
    if not doc_ref:
        return {"status": "error", "message": "Firestore not available"}
    try:
        doc_ref.set(
            {
                "name": name,
                "project": project_raw,
                "created_at": firestore.SERVER_TIMESTAMP,
                "updated_at": firestore.SERVER_TIMESTAMP,
            }
        )
        return {
            "project_id": project_id,
            "project_raw": project_raw,
            "name": name,
        }
    except Exception as exc:
        logger.warning(f"Failed to create project: {exc}")
        return {"status": "error", "message": "Failed to create project"}


@app.get("/projects/{project_id}")
async def get_project(project_id: str) -> Dict[str, object]:
    doc_ref = get_project_doc_ref(project_id)
    if not doc_ref:
        return {"status": "error", "message": "Firestore not available"}
    try:
        snapshot = doc_ref.get()
        if not snapshot.exists:
            return {"status": "not_found"}
        data = snapshot.to_dict() or {}
        try:
            project_raw = data.get("project")
            logger.info(
                "Load project: project_id=%s has_project=%s size_bytes=%s",
                project_id,
                isinstance(project_raw, dict),
                estimate_project_bytes(project_raw) if isinstance(project_raw, dict) else None,
            )
        except Exception:
            pass
        return {
            "project_raw": data.get("project") or None,
            "project": normalize_project_state(data.get("project") or {}).model_dump(),
            "name": data.get("name") or "Untitled project",
        }
    except Exception as exc:
        logger.warning(f"Failed to load project: {exc}")
        return {"status": "error", "message": "Failed to load project"}


@app.post("/projects/{project_id}/duplicate")
async def duplicate_project(project_id: str, payload: Dict = Body(default={})) -> Dict[str, object]:
    client = get_firestore_client()
    if not client:
        return {"status": "error", "message": "Firestore not available"}
    source_ref = get_project_doc_ref(project_id)
    if not source_ref:
        return {"status": "error", "message": "Firestore not available"}
    try:
        snapshot = source_ref.get()
        if not snapshot.exists:
            return {"status": "not_found"}
        data = snapshot.to_dict() or {}
        project_raw = data.get("project")
        if not isinstance(project_raw, dict):
            project_raw = {}
        limit_error = validate_project_limits(project_raw)
        if limit_error:
            return {
                "status": "error",
                **limit_error,
            }
        source_name = data.get("name") or "Untitled project"
        name = str(payload.get("name") or f"Kopia av {source_name}").strip() or f"Kopia av {source_name}"
        project_id = str(uuid4())
        doc_ref = get_project_doc_ref(project_id)
        if not doc_ref:
            return {"status": "error", "message": "Firestore not available"}
        doc_ref.set(
            {
                "name": name,
                "project": project_raw,
                "created_at": firestore.SERVER_TIMESTAMP,
                "updated_at": firestore.SERVER_TIMESTAMP,
            }
        )
        return {
            "project_id": project_id,
            "project_raw": project_raw,
            "name": name,
        }
    except Exception as exc:
        logger.warning(f"Failed to duplicate project: {exc}")
        return {"status": "error", "message": "Failed to duplicate project"}


@app.patch("/projects/{project_id}")
async def rename_project(project_id: str, payload: Dict = Body(...)) -> Dict[str, str]:
    doc_ref = get_project_doc_ref(project_id)
    if not doc_ref:
        return {"status": "error", "message": "Firestore not available"}
    name = str(payload.get("name") or "").strip()
    if not name:
        return {"status": "invalid", "message": "Project name is required"}
    try:
        doc_ref.set(
            {
                "name": name,
                "updated_at": firestore.SERVER_TIMESTAMP,
            },
            merge=True,
        )
        return {"status": "ok", "name": name}
    except Exception as exc:
        logger.warning(f"Failed to rename project: {exc}")
        return {"status": "error", "message": "Failed to rename project"}


@app.delete("/projects/{project_id}")
async def delete_project(project_id: str, request: Request) -> Dict[str, str]:
    require_admin(request)
    global last_saved_project_hash
    doc_ref = get_project_doc_ref(project_id)
    if not doc_ref:
        return {"status": "error", "message": "Firestore not available"}
    try:
        doc_ref.delete()
        last_saved_project_hash.pop(project_id, None)
        return {"status": "ok"}
    except Exception as exc:
        logger.warning(f"Failed to delete project: {exc}")
        return {"status": "error", "message": "Failed to delete project"}


@app.post("/projects/{project_id}/state")
async def set_project_state_for_project(project_id: str, payload: Dict = Body(...)) -> Dict[str, str]:
    try:
        project_raw = payload.get("project_raw") or payload.get("project") or payload
        project_name = payload.get("name")
        force_overwrite = bool(payload.get("force"))
        if not isinstance(project_raw, dict):
            return {"status": "invalid"}

        # Canonicalize updated_at using server time to avoid client clock skew.
        prev_updated_at = project_raw.get("updated_at")
        prev_ms = prev_updated_at if isinstance(prev_updated_at, int) else 0
        project_raw["updated_at"] = monotonic_updated_at_ms(prev_ms)

        try:
            doc_list = project_raw.get("documents", [])
            docs_count = len(doc_list) if isinstance(doc_list, list) else 0
            sample = doc_list[0] if docs_count and isinstance(doc_list[0], dict) else {}
            logger.info(
                "Save state: project_id=%s docs=%s sample_keys=%s size_bytes=%s",
                project_id,
                docs_count,
                list(sample.keys())[:8] if isinstance(sample, dict) else [],
                estimate_project_bytes(project_raw),
            )
        except Exception:
            # Keep saves working even if diagnostics fail.
            pass

        limit_error = validate_project_limits(project_raw)
        if limit_error:
            return {
                "status": "error",
                **limit_error,
            }

        # Safety net: avoid wiping a real project due to a client race/bug.
        # Blocks both obviously empty payloads and suspiciously small payloads compared to existing state.
        if not force_overwrite:
            incoming_bytes = estimate_project_bytes(project_raw) or 0
            existing = load_project_from_firestore(project_id)
            existing_bytes = estimate_project_bytes(existing) if isinstance(existing, dict) else 0

            existing_doc_chars = estimate_document_content_chars(existing) if isinstance(existing, dict) else 0
            incoming_doc_chars = estimate_document_content_chars(project_raw)

            looks_like_content_wipe = (
                isinstance(existing, dict)
                and existing_doc_chars > 0
                and incoming_doc_chars == 0
            )

            looks_empty = is_project_effectively_empty(project_raw) or incoming_bytes <= 32
            looks_suspiciously_small = (
                isinstance(existing, dict)
                and existing_bytes is not None
                and incoming_bytes is not None
                and existing_bytes >= NONEMPTY_PROJECT_MIN_BYTES
                and incoming_bytes <= EMPTY_OVERWRITE_MAX_BYTES
            )

            if looks_empty or looks_suspiciously_small:
                if isinstance(existing, dict) and not is_project_effectively_empty(existing):
                    logger.warning(
                        "Refusing overwrite. project_id=%s incoming_bytes=%s existing_bytes=%s empty=%s suspicious=%s",
                        project_id,
                        incoming_bytes,
                        existing_bytes,
                        looks_empty,
                        looks_suspiciously_small,
                    )
                    return JSONResponse(
                        status_code=409,
                        content={
                            "status": "error",
                            "reason": "empty_overwrite",
                            "message": "Refusing to overwrite a non-empty project with an empty or suspiciously small payload.",
                        },
                    )

            # Additional guard: prevent saving a payload that wipes all document content.
            if looks_like_content_wipe:
                logger.warning(
                    "Refusing content wipe. project_id=%s incoming_doc_chars=%s existing_doc_chars=%s incoming_bytes=%s existing_bytes=%s",
                    project_id,
                    incoming_doc_chars,
                    existing_doc_chars,
                    incoming_bytes,
                    existing_bytes,
                )
                return JSONResponse(
                    status_code=409,
                    content={
                        "status": "error",
                        "reason": "content_wipe",
                        "message": "Refusing to overwrite a project with a payload that contains no document content.",
                    },
                )

        set_in_memory_project(project_id, project_raw)
        saved_ok = save_project_to_firestore(project_raw, project_id=project_id)
        if not saved_ok:
            return JSONResponse(
                status_code=503,
                content={
                    "status": "error",
                    "message": "Failed to save project to Firestore",
                },
            )
        if project_name:
            doc_ref = get_project_doc_ref(project_id)
            if doc_ref:
                doc_ref.set({"name": project_name}, merge=True)
        return {"status": "ok", "updated_at": project_raw.get("updated_at")}
    except Exception:
        logger.exception("Failed to save project state")
        return {"status": "error"}


@app.post("/project/state")
async def set_project_state(payload: Dict = Body(...)) -> Dict[str, str]:
    try:
        project_raw = payload.get("project_raw") or payload.get("project") or payload
        if not isinstance(project_raw, dict):
            return {"status": "invalid"}

        prev_updated_at = project_raw.get("updated_at")
        prev_ms = prev_updated_at if isinstance(prev_updated_at, int) else 0
        project_raw["updated_at"] = monotonic_updated_at_ms(prev_ms)

        limit_error = validate_project_limits(project_raw)
        if limit_error:
            return {
                "status": "error",
                **limit_error,
            }

        project_id = DEFAULT_PROJECT_ID
        set_in_memory_project(project_id, project_raw)
        saved_ok = save_project_to_firestore(project_raw)
        if not saved_ok:
            return JSONResponse(
                status_code=503,
                content={
                    "status": "error",
                    "message": "Failed to save project to Firestore",
                },
            )
        await manager.broadcast(
            project_id,
            {
                "type": "project:update",
                "project": get_in_memory_project(project_id).model_dump(),
                "project_raw": get_in_memory_project_raw(project_id),
            },
        )
        return {"status": "ok", "updated_at": project_raw.get("updated_at")}
    except Exception as exc:
        logger.exception("Failed to save project state")
        return {"status": "error"}


@app.post("/project/load")
async def load_project(file: UploadFile = File(...)) -> Dict[str, str]:
    payload = await file.read()
    project_state = ProjectState.model_validate_json(payload)

    project_id = DEFAULT_PROJECT_ID
    project_raw = project_state.model_dump()
    limit_error = validate_project_limits(project_raw)
    if limit_error:
        return {
            "status": "error",
            **limit_error,
        }
    set_in_memory_project(project_id, project_raw)
    saved_ok = save_project_to_firestore(project_raw)
    if not saved_ok:
        return JSONResponse(
            status_code=503,
            content={
                "status": "error",
                "message": "Failed to save project to Firestore",
            },
        )
    await manager.broadcast(project_id, get_in_memory_project(project_id).model_dump())

    return {"status": "ok"}


@app.get("/project/save")
async def save_project() -> Response:
    content = get_in_memory_project(DEFAULT_PROJECT_ID).model_dump_json(indent=2)
    headers = {
        "Content-Disposition": "attachment; filename=project_backup.json",
    }
    return Response(content=content, media_type="application/json", headers=headers)


@app.get("/export/word")
async def export_word(project_id: Optional[str] = None) -> Response:
    resolved_id = resolve_project_id(project_id)
    project_raw = load_project_from_firestore(resolved_id) if project_id else None
    project = normalize_project_state(project_raw) if project_raw else get_in_memory_project(resolved_id)
    document = Document()
    document.add_heading("Grounded Theory Analysrapport", level=1)

    document.add_heading("Teori (Selektiv kodning)", level=2)
    def get_code_by_id_local(code_id: str) -> Optional[CodeItem]:
        return next((code for code in project.codes if code.id == code_id), None)

    def get_category_by_id_local(category_id: str) -> Optional[CategoryItem]:
        return next((category for category in project.categories if category.id == category_id), None)

    def get_document_by_id_local(document_id: str) -> Optional[DocumentItem]:
        return next((doc for doc in project.documents if doc.id == document_id), None)

    core_category = (
        get_category_by_id_local(project.core_category_id)
        if project.core_category_id
        else None
    )
    document.add_paragraph(
        f"Kärnkategori: {core_category.name if core_category else 'Inte vald'}"
    )
    document.add_paragraph(
        project.theory_description or "Ingen teoribeskrivning angiven."
    )

    document.add_heading("Kategorier (Axial kodning)", level=2)
    for category in project.categories:
        document.add_paragraph(category.name)
        if category.precondition or category.action or category.consequence:
            if category.precondition:
                document.add_paragraph(f"Förutsättning: {category.precondition}")
            if category.action:
                document.add_paragraph(f"Handling: {category.action}")
            if category.consequence:
                document.add_paragraph(f"Konsekvens: {category.consequence}")
        for code_id in category.contained_code_ids:
            code = get_code_by_id_local(code_id)
            if code:
                document.add_paragraph(f"- {code.name}")

    document.add_heading("Evidens (Öppen kodning)", level=2)
    for code in project.codes:
        document.add_paragraph(code.name)
        related_highlights = [highlight for highlight in project.highlights if highlight.code_id == code.id]
        if not related_highlights:
            document.add_paragraph("Inga citat ännu.")
            continue
        for highlight in related_highlights:
            document_item = get_document_by_id_local(highlight.document_id)
            if not document_item:
                continue
            quote = document_item.content[highlight.start_index : highlight.end_index]
            document.add_paragraph(f'"{quote}"')

    output = BytesIO()
    document.save(output)
    output.seek(0)

    headers = {
        "Content-Disposition": "attachment; filename=grounded_theory_report.docx",
    }
    return Response(
        content=output.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.get("/export/excel")
async def export_excel(project_id: Optional[str] = None) -> Response:
    resolved_id = resolve_project_id(project_id)
    project_raw = load_project_from_firestore(resolved_id) if project_id else None
    project = normalize_project_state(project_raw) if project_raw else get_in_memory_project(resolved_id)
    rows = []
    code_by_id = {code.id: code for code in project.codes}
    category_by_id = {category.id: category for category in project.categories}
    categories_by_code: Dict[str, List[str]] = {}
    for category in project.categories:
        for code_id in category.contained_code_ids:
            categories_by_code.setdefault(code_id, []).append(category.name)

    for highlight in project.highlights:
        code = code_by_id.get(highlight.code_id)
        document_item = next((doc for doc in project.documents if doc.id == highlight.document_id), None)
        if not code or not document_item:
            continue
        quote = document_item.content[highlight.start_index : highlight.end_index]
        rows.append(
            {
                "Code": code.name,
                "Category": ", ".join(categories_by_code.get(code.id, [])),
                "Förutsättning": "; ".join(
                    filter(
                        None,
                        [
                            category_by_id[category_id].precondition
                            for category_id in category_by_id
                            if code.id in category_by_id[category_id].contained_code_ids
                        ],
                    )
                ),
                "Handling": "; ".join(
                    filter(
                        None,
                        [
                            category_by_id[category_id].action
                            for category_id in category_by_id
                            if code.id in category_by_id[category_id].contained_code_ids
                        ],
                    )
                ),
                "Konsekvens": "; ".join(
                    filter(
                        None,
                        [
                            category_by_id[category_id].consequence
                            for category_id in category_by_id
                            if code.id in category_by_id[category_id].contained_code_ids
                        ],
                    )
                ),
                "Dokument": document_item.title,
                "Citat": quote,
            }
        )

    df = pd.DataFrame(
        rows
        or [
            {
                "Code": "",
                "Category": "",
                "Förutsättning": "",
                "Handling": "",
                "Konsekvens": "",
                "Dokument": "",
                "Citat": "",
            }
        ]
    )
    output = BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Grounded Theory")
    output.seek(0)

    headers = {
        "Content-Disposition": "attachment; filename=grounded_theory_report.xlsx",
    }
    return Response(
        content=output.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )


# Catch-all route for SPA - must be defined LAST
@app.get("/{full_path:path}")
async def serve_spa(full_path: str):
    # If the file exists in dist (e.g. favicon.ico, manifest.json), serve it
    possible_file = FRONTEND_DIST / full_path
    if possible_file.is_file():
        return FileResponse(possible_file)

    # Otherwise serve index.html for client-side routing
    index_path = FRONTEND_DIST / "index.html"
    if index_path.exists():
        return FileResponse(index_path)
    
    return Response("Frontend not found. Did you run 'npm run build'?", status_code=404)
